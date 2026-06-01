import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_md(text):
    # Remove markdown bold/italic/latex math dollar signs
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    # Remove markdown link markup like [link text](url) -> link text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text.strip()

def read_markdown_sections(file_path):
    sections = {}
    current_section = None
    current_content = []
    
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        match = re.match(r'^(##|###)\s+(.*)', line)
        if match:
            if current_section:
                sections[current_section] = "\n".join(current_content).strip()
            # Clean section name and keep it lowercase to map easily
            current_section = match.group(2).strip().lower()
            # clean tags or numbers
            current_section = re.sub(r'^[\d\.]+\s+', '', current_section)
            current_content = []
        else:
            if current_section is not None:
                current_content.append(line.rstrip())
                
    if current_section:
        sections[current_section] = "\n".join(current_content).strip()
        
    return sections

def inject_to_docx():
    docx_path = "Tesis de Investigación 2026 Avance capitulo 1-2docx.docx"
    doc = Document(docx_path)
    
    # Read Markdown sections
    cap1 = read_markdown_sections("docs/10-capitulo1.md")
    cap2_ant = read_markdown_sections("docs/20-capitulo2-antecedentes.md")
    cap2_art = read_markdown_sections("docs/21-capitulo2-estadoarte.md")
    cap2_teo = read_markdown_sections("docs/22-capitulo2-marcoteorico.md")
    
    # Combine sections map
    sec_map = {}
    # Combine all
    for k, v in cap1.items():
        sec_map[k] = v
    for k, v in cap2_ant.items():
        sec_map[k] = v
    for k, v in cap2_art.items():
        sec_map[k] = v
    for k, v in cap2_teo.items():
        sec_map[k] = v
        
    print("Mapped Markdown Sections:", list(sec_map.keys()))
    
    # We will locate paragraph by text match
    for i, p in enumerate(doc.paragraphs):
        p_text_clean = p.text.strip().lower()
        # Clean section number prefixes if any
        p_text_normalized = re.sub(r'^[\d\.]+\s+', '', p_text_clean)
        
        # Exact or fuzzy match of headings in docx to our markdown keys
        matched_key = None
        for key in sec_map.keys():
            if key in p_text_normalized or p_text_normalized in key:
                if len(p_text_normalized) > 5: # avoid tiny heading matches
                    matched_key = key
                    break
        
        if matched_key:
            print(f"Matched Docx: '{p.text}' with Key: '{matched_key}'")
            # If we match, we will append our text after this heading.
            # But let's check if the next paragraphs are blank or have placeholders and delete them.
            # For simplicity, we can insert our paragraphs after this paragraph.
            content_text = sec_map[matched_key]
            paragraphs_to_add = content_text.split('\n\n')
            
            # Insert paragraphs
            current_p = p
            for block in paragraphs_to_add:
                block_clean = clean_md(block)
                if not block_clean:
                    continue
                # Add paragraph
                new_p = doc.add_paragraph()
                # Insert after current_p (requires editing document XML element order)
                p._element.addnext(new_p._element)
                new_p.text = block_clean
                # Set Times New Roman, 12pt, double space
                for run in new_p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                # update current_p
                current_p = new_p
                
    # Update Cronograma to insert Gantt image
    for p in doc.paragraphs:
        if "cronograma" in p.text.strip().lower():
            print("Found Cronograma heading, inserting Gantt chart details...")
            # Insert Gantt chart details
            phases_desc = (
                "Para asegurar la excelencia académica y un desarrollo de software robusto con sus respectivas pruebas estadísticas y de usabilidad, el cronograma de la tesis se ha reestructurado para durar desde Mayo de 2026 hasta la primera semana de Diciembre de 2026.\n\n"
                "- Fase 1: Preparación y Tratamiento de Datos (Mayo 2026). Definición de variables, recopilación de datos de fuentes públicas (MIDAGRI, SENAMHI, SENASA, SUNAT) e inyección de anomalías del dataset sintético.\n"
                "- Fase 2: Desarrollo Backend y Modelado (Junio - Julio 2026). Configuración del backend en Flask, entrenamiento de GBDTs (XGBoost, LightGBM, CatBoost) y ensembles de detección de anomalías (Isolation Forest, LOF, ECOD) con PyOD.\n"
                "- Fase 3: Explicabilidad y Reportes RAG (Agosto 2026). Cálculo automatizado de valores SHAP a nivel de instancia y estructuración del flujo RAG con prompts controlados para generar explicaciones narrativas sin alucinaciones.\n"
                "- Fase 4: Integración del Pipeline y UI Dashboard (Septiembre 2026). Orquestación del flujo completo y desarrollo del panel de control web con visualizaciones interactivas de alertas y SHAP.\n"
                "- Fase 5: Protocolo de Usabilidad con Testers (Octubre 2026). Ejecución de pruebas con 10 o más analistas operativos, registrando tiempos de respuesta, nivel de comprensión y usabilidad (SUS).\n"
                "- Fase 6: Pruebas de Calidad e Implementación de Cambios (Noviembre 2026). Ajuste del sistema a partir de las pruebas de usabilidad e integración de retroalimentación de los usuarios.\n"
                "- Fase 7: Redacción Final de Capítulos y Anexos (Noviembre 2026). Consolidación del marco metodológico, redacción de capítulos de resultados y análisis del impacto de la explicabilidad.\n"
                "- Fase 8: Revisiones Finales y Sustentación (Diciembre 2026). Compilación final del documento, levantamiento de observaciones del jurado académico y defensa oral de la tesis."
            )
            
            current_p = p
            for block in phases_desc.split('\n\n'):
                new_p = doc.add_paragraph()
                current_p._element.addnext(new_p._element)
                new_p.text = block
                for run in new_p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                current_p = new_p
            
            # Insert Gantt chart image
            gantt_img_path = "data/downloads/gantt_chart.png"
            if os.path.exists(gantt_img_path):
                img_p = doc.add_paragraph()
                current_p._element.addnext(img_p._element)
                run = img_p.add_run()
                run.add_picture(gantt_img_path, width=Inches(6.0))
                # Add label under image
                lbl_p = doc.add_paragraph()
                img_p._element.addnext(lbl_p._element)
                lbl_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_lbl = lbl_p.add_run("Figura 1.1: Cronograma Visual del Proyecto (Diagrama de Gantt)")
                run_lbl.font.name = 'Times New Roman'
                run_lbl.font.size = Pt(10)
                run_lbl.italic = True
                
    # Update Table 0 (Técnicas e Instrumentos)
    if len(doc.tables) > 0:
        table_0 = doc.tables[0]
        # Populate Table 0 with:
        # Fila 1: Revisión sistemática de literatura | Fichas de lectura / Gestor de referencias (Zotero) | Mapear antecedentes, estado del arte y sustentar decisiones de diseño experimental
        # Fila 2: Experimentación computacional (E1-E5) | Scripts de Python / TensorBoard / Model Cards / Datasheets | Evaluar rendimiento de detección (PR-AUC, F1), SHAP (cobertura) y RAG (RÚBRICA)
        # Fila 3: Evaluación con usuarios (Usabilidad) | Escala SUS / Protocolo de tareas cronometradas / Likert | Medir tiempo de decisión, nivel de comprensión y usabilidad del panel integrado
        data_rows = [
            ("Revisión sistemática de literatura", "Fichas de lectura y gestor bibliográfico (Zotero)", "Sustentar el estado del arte y contrastar el gap de investigación identificado"),
            ("Experimentación controlada (E1-E5)", "Módulos de telemetría, scripts de Python, y rúbricas de evaluación", "Evaluar métricas técnicas (PR-AUC, F1-Score, consistencia RAG y cobertura SHAP)"),
            ("Evaluación de usabilidad con analistas", "Cuestionarios Likert, escala SUS, y registro cronometrado de tiempo-a-decisión", "Medir el impacto de la explicabilidad en la toma de decisiones operativas")
        ]
        
        # Table 0 columns: Técnica | Instrumento | Propósito
        # Row 0 is header
        for row_idx, row_data in enumerate(data_rows):
            if row_idx + 1 < len(table_0.rows):
                row = table_0.rows[row_idx + 1]
            else:
                row = table_0.add_row()
            for col_idx, text in enumerate(row_data):
                row.cells[col_idx].text = text
                # font formatting
                for p in row.cells[col_idx].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)

    out_docx_path = "Tesis de Investigación 2026 Avance capitulo 1-2docx_Actualizado.docx"
    doc.save(out_docx_path)
    print(f"Updated document saved to: {out_docx_path}")

if __name__ == '__main__':
    inject_to_docx()
