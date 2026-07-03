import docx
from pathlib import Path

docx_path = Path("Tesis de Investigación YOSET.docx")
out_path = Path("output/tesis_yoset_content.txt")

if not docx_path.exists():
    # Try with normal 'o' instead of combining Unicode if applicable
    docx_path = Path("Tesis de Investigación YOSET.docx")

print(f"Reading from: {docx_path.absolute()}")
doc = docx.Document(docx_path)
lines = []

lines.append("=========================================")
lines.append(f"Documento: {docx_path.name}")
lines.append(f"Total parrafos: {len(doc.paragraphs)}")
lines.append(f"Total tablas: {len(doc.tables)}")
lines.append("=========================================\n")

# Extract paragraphs
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        lines.append(f"[P{i}] {text}")

# Extract tables
for t_idx, table in enumerate(doc.tables):
    lines.append(f"\n--- TABLA {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        try:
            row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(f"  Row {r_idx}: " + " | ".join(row_cells))
        except Exception as e:
            lines.append(f"  Row {r_idx} Error: {str(e)}")

with open(out_path, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Contenido guardado exitosamente en: {out_path.absolute()}")
