import docx
from pathlib import Path

docx_path = Path("avance/Plantilla - Tesis de Investigación 2026 (1).docx")
out_path = Path("data/downloads/docx_content.txt")

doc = docx.Document(docx_path)
lines = []

lines.append("=========================================")
print("Abierto docx con exito.")
print(f"Total parrafos: {len(doc.paragraphs)}")
print(f"Total tablas: {len(doc.tables)}")

# Extract paragraphs
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        lines.append(f"[P{i}] {text}")

# Extract tables
for t_idx, table in enumerate(doc.tables):
    lines.append(f"\n--- TABLA {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        lines.append(f"  Row {r_idx}: " + " | ".join(row_cells))

with open(out_path, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Contenido guardado en {out_path}")
